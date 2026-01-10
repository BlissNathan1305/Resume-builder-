import streamlit as st
from reportlab.lib.pagesizes import LETTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, RGBColor
import io

st.set_page_config(page_title="Premium CV Builder", layout="wide")
st.title("Premium CV / Resume Builder")

# -----------------------
# Theme & Styling
# -----------------------
st.sidebar.header("Theme & Styling")
template = st.sidebar.selectbox("Choose Template", ["Classic ATS", "Modern Professional", "Designer Minimal"])
accent_color = st.sidebar.selectbox("Accent Color", ["Blue", "Gray", "Green", "Black"])

# -----------------------
# Personal Information
# -----------------------
st.header("Personal Information")
name = st.text_input("Full Name")
title = st.text_input("Professional Title")
email = st.text_input("Email")
phone = st.text_input("Phone")
location = st.text_input("Location")
linkedin = st.text_input("LinkedIn URL")
github = st.text_input("GitHub URL")
portfolio = st.text_input("Portfolio URL")

# -----------------------
# Professional Summary
# -----------------------
st.header("Professional Summary")
summary_input = st.text_area("Write your professional summary here", height=100)

# -----------------------
# Dynamic Sections
# -----------------------
def add_dynamic_section(section_name, fields):
    st.header(section_name)
    items = []
    num_items = st.number_input(f"How many {section_name} entries?", min_value=1, max_value=20, value=1, step=1)
    for i in range(num_items):
        st.subheader(f"{section_name} #{i+1}")
        item = {}
        for field in fields:
            item[field] = st.text_input(f"{field} (Entry #{i+1})")
        items.append(item)
    return items

work_fields = ["Job Title", "Company", "Start Date - End Date", "Responsibilities (comma-separated)"]
education_fields = ["Degree", "Institution", "Start Date - End Date", "Details"]
project_fields = ["Project Name", "Description", "Technologies Used"]

work_experience = add_dynamic_section("Work Experience", work_fields)
education = add_dynamic_section("Education", education_fields)
projects = add_dynamic_section("Projects", project_fields)
skills = st.text_input("Skills (comma-separated)")
certifications = st.text_input("Certifications (comma-separated)")

# -----------------------
# PDF Generation Function
# -----------------------
def generate_pdf():
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=LETTER, rightMargin=72,leftMargin=72,topMargin=72,bottomMargin=18)
    styles = getSampleStyleSheet()
    story = []

    # Template styles
    if template == "Classic ATS":
        heading_color = colors.black
        heading_style = ParagraphStyle('HeadingStyle', parent=styles['Heading2'], textColor=heading_color, fontSize=14, spaceAfter=6)
        subheading_style = ParagraphStyle('SubHeading', parent=styles['Heading3'], textColor=heading_color, fontSize=12, spaceAfter=4)
        normal_style = styles['Normal']
    elif template == "Modern Professional":
        color_map = {"Blue": colors.HexColor("#0b5394"),
                     "Gray": colors.HexColor("#4d4d4d"),
                     "Green": colors.HexColor("#38761d"),
                     "Black": colors.black}
        heading_color = color_map.get(accent_color, colors.HexColor("#0b5394"))
        heading_style = ParagraphStyle('HeadingStyle', parent=styles['Heading2'], textColor=heading_color, fontSize=16, spaceAfter=8)
        subheading_style = ParagraphStyle('SubHeading', parent=styles['Heading3'], textColor=heading_color, fontSize=13, spaceAfter=6)
        normal_style = styles['Normal']
    else:  # Designer Minimal
        heading_color = colors.HexColor("#222222")
        heading_style = ParagraphStyle('HeadingStyle', parent=styles['Heading2'], textColor=heading_color, fontSize=16, spaceAfter=8, leading=18)
        subheading_style = ParagraphStyle('SubHeading', parent=styles['Heading3'], textColor=heading_color, fontSize=13, spaceAfter=6, leading=15)
        normal_style = styles['Normal']

    # Personal Info
    story.append(Paragraph(f"<b>{name}</b>", heading_style))
    story.append(Paragraph(f"{title}", normal_style))
    story.append(Paragraph(f"{email} | {phone} | {location}", normal_style))
    links = " | ".join([linkedin, github, portfolio])
    if links.strip() != "|":
        story.append(Paragraph(links, normal_style))
    story.append(Spacer(1, 12))

    # Professional Summary
    story.append(Paragraph("Professional Summary", heading_style))
    story.append(Paragraph(summary_input, normal_style))
    story.append(Spacer(1, 12))

    # Work Experience
    story.append(Paragraph("Work Experience", heading_style))
    for job in work_experience:
        story.append(Paragraph(f"{job['Job Title']} | {job['Company']} | {job['Start Date - End Date']}", subheading_style))
        bullets = job['Responsibilities (comma-separated)'].split(',')
        for b in bullets:
            story.append(Paragraph(f"• {b.strip()}", normal_style))
        story.append(Spacer(1, 6))

    # Education
    story.append(Paragraph("Education", heading_style))
    for edu in education:
        story.append(Paragraph(f"{edu['Degree']} | {edu['Institution']} | {edu['Start Date - End Date']}", subheading_style))
        story.append(Paragraph(edu['Details'], normal_style))
        story.append(Spacer(1, 6))

    # Projects
    if any([p["Project Name"] for p in projects]):
        story.append(Paragraph("Projects", heading_style))
        for proj in projects:
            story.append(Paragraph(proj['Project Name'], subheading_style))
            story.append(Paragraph(proj['Description'], normal_style))
            story.append(Paragraph(f"Technologies Used: {proj['Technologies Used']}", normal_style))
            story.append(Spacer(1, 6))

    # Skills
    story.append(Paragraph("Skills", heading_style))
    story.append(Paragraph(skills, normal_style))
    story.append(Spacer(1, 6))

    # Certifications
    if certifications.strip() != "":
        story.append(Paragraph("Certifications", heading_style))
        story.append(Paragraph(certifications, normal_style))
        story.append(Spacer(1, 6))

    doc.build(story)
    pdf = buffer.getvalue()
    buffer.close()
    return pdf

# -----------------------
# DOCX Generation Function
# -----------------------
def generate_docx():
    doc = Document()
    heading_size = Pt(16)
    subheading_size = Pt(13)

    def add_heading(text, level):
        h = doc.add_heading(text, level=level)
        run = h.runs[0]
        if level == 1:
            run.font.size = heading_size
        else:
            run.font.size = subheading_size
        return h

    add_heading(name, 0)
    doc.add_paragraph(title)
    doc.add_paragraph(f"{email} | {phone} | {location}")
    links = " | ".join([linkedin, github, portfolio])
    if links.strip() != "|":
        doc.add_paragraph(links)

    add_heading("Professional Summary", 1)
    doc.add_paragraph(summary_input)

    add_heading("Work Experience", 1)
    for job in work_experience:
        add_heading(f"{job['Job Title']} | {job['Company']} | {job['Start Date - End Date']}", 2)
        bullets = job['Responsibilities (comma-separated)'].split(',')
        for b in bullets:
            doc.add_paragraph(b.strip(), style='ListBullet')

    add_heading("Education", 1)
    for edu in education:
        add_heading(f"{edu['Degree']} | {edu['Institution']} | {edu['Start Date - End Date']}", 2)
        doc.add_paragraph(edu['Details'])

    if any([p["Project Name"] for p in projects]):
        add_heading("Projects", 1)
        for proj in projects:
            add_heading(proj['Project Name'], 2)
            doc.add_paragraph(proj['Description'])
            doc.add_paragraph(f"Technologies Used: {proj['Technologies Used']}")

    add_heading("Skills", 1)
    doc.add_paragraph(skills)

    if certifications.strip() != "":
        add_heading("Certifications", 1)
        doc.add_paragraph(certifications)

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    return docx_bytes

# -----------------------
# Download Buttons
# -----------------------
col1, col2 = st.columns(2)
with col1:
    if st.button("Download PDF"):
        pdf = generate_pdf()
        st.download_button(label="Download PDF", data=pdf, file_name=f"{name}_CV.pdf", mime="application/pdf")
with col2:
    if st.button("Download DOCX"):
        docx_bytes = generate_docx()
        st.download_button(label="Download DOCX", data=docx_bytes, file_name=f"{name}_CV.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
